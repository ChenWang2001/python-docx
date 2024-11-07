import docx
import pandas as pd
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT, WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement

from docx.oxml.ns import qn

# 用于设置的word中表格单元格的背景颜色设置
def set_cell_color(target_cell, color='5B9BD5'):
    # 创建 tcPr 元素
    tcPr = target_cell._element.tcPr if target_cell._element.tcPr is not None else OxmlElement('w:tcPr')

    # 创建 shd 元素并设置背景颜色
    shd = OxmlElement('w:shd')
    shd.set(qn('w:fill'), color)  # 设置背景颜色为红色（FF0000）

    # 将 shd 元素添加到 tcPr 元素中
    tcPr.append(shd)

    # 将 tcPr 元素添加到单元格中
    if target_cell._element.tcPr is None:
        target_cell._element.append(tcPr)

# 添加单元格后将原始样式复制给新生成的单元格
def copy_font_style(source_cell, target_cell):
    # 获取源单元格（第一行第一列）
    source_paragraph = source_cell.paragraphs[0]
    source_run = source_paragraph.runs[0]

    target_paragraph = target_cell.paragraphs[0] if target_cell.paragraphs else target_cell.add_paragraph()
    target_run = target_paragraph.runs[0]

    # 复制段落格式
    target_paragraph.style = source_paragraph.style
    target_paragraph.paragraph_format.alignment = source_paragraph.paragraph_format.alignment
    target_paragraph.paragraph_format.line_spacing = source_paragraph.paragraph_format.line_spacing
    target_paragraph.paragraph_format.space_before = source_paragraph.paragraph_format.space_before
    target_paragraph.paragraph_format.space_after = source_paragraph.paragraph_format.space_after

    # 复制字体格式
    target_run.bold = source_run.bold
    target_run.italic = source_run.italic
    target_run.underline = source_run.underline
    target_run.font.size = source_run.font.size

